from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy

SRC = 'efp_modernization_work/modernization-plan-efp.updated.docx'
OUT = 'efp_modernization_work/modernization-plan-efp.daily-reporting.docx'

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_text(cell, text, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

def add_row(table, values):
    cells = table.add_row().cells
    for cell, value in zip(cells, values):
        set_cell_text(cell, value)
    return cells

def insert_paragraph_after(paragraph, text=None, style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    p = paragraph._parent.add_paragraph()
    p._p = new_p
    p._element = new_p
    if style:
        p.style = style
    if text:
        p.add_run(text)
    return p


doc = Document(SRC)

# Update roadmap narrative to call out the two field-data tools explicitly.
old = (
    'Narrative summary. Pre-close work is dominated by visible-state compliance and transaction-mechanics items: new-hire form cleanup, UIA reconstruction, office succession, BFS certification verification, and the sunset of the legacy banking signature-stamp practice. The first 90 days focus on business-continuity risks: inspection-delivery concentration, field-to-office communication, IT backup/security, safety program management, and compliance tracking. Year 1 shifts toward structural modernization: QBO/payroll, estimating documentation, inspection-service strategy, QA/QC, project-management discipline, customer/revenue segmentation, MDM, fleet and inventory hygiene, and management reporting. Strategic items should not be forced into the first 90 days unless they become prerequisites for revenue protection or buyer/lender requirements.'
)
new = (
    'Narrative summary. Pre-close work is dominated by visible-state compliance and transaction-mechanics items: new-hire form cleanup, UIA reconstruction, office succession, BFS certification verification, and the sunset of the legacy banking signature-stamp practice. The first 90 days focus on business-continuity risks and field-data capture: inspection-delivery concentration, field-to-office communication, daily job site reporting/photo documentation, digital timecard submission, IT backup/security, safety program management, and compliance tracking. Year 1 shifts toward structural modernization: QBO/payroll, estimating documentation, inspection-service strategy, QA/QC, project-management discipline, customer/revenue segmentation, MDM, fleet and inventory hygiene, and management reporting. Strategic items should not be forced into the first 90 days unless they become prerequisites for revenue protection or buyer/lender requirements.'
)
for p in doc.paragraphs:
    if p.text == old:
        p.text = new
        break

# Add dependency/critical-path rows.
dep = doc.tables[1]
add_row(dep, [
    'Daily job site reporting / photo documentation',
    'Creates daily evidence for progress, blockers, safety conditions, change-order support, QA/QC checks, customer updates, and billing readiness. This should be treated as a Project Delivery & Operations control, not just a communications convenience.',
])
add_row(dep, [
    'Digital timecard submission / labor cost capture',
    'Feeds payroll, certified payroll, job-costing, WIP timing, and revenue recognition. Phase 1 can be photo/email/text submission of handwritten sheets; later phases can move to QBO Time, Workyard, busybusy, ExakTime, or another field-labor platform once union payroll and job-costing requirements are known.',
])

# Add/update workflow backlog rows.
wf = doc.tables[2]
add_row(wf, [
    'Daily Job Site Reporting / Photo Documentation',
    'Daily foreman or crew-lead update by job: photos, manpower, percent-complete signal, material constraints, inspections/AHJ readiness, safety incidents/near misses, blockers, change-order indicators, and work completed since prior report.',
    'Supports Findings #17, #19, #20, #22, and #23; candidate tools include CompanyCam or an equivalent photo/reporting workflow. Provides evidence for project tracking, QA/QC, change orders, customer communication, and billing readiness.',
])
# Replace payroll/timecard row with more explicit phased approach.
for row in wf.rows:
    if row.cells[0].text.strip() == 'Payroll / Timecard Flow':
        set_cell_text(row.cells[0], 'Payroll / Digital Timecard Submission Flow')
        set_cell_text(row.cells[1], 'Current paper timesheets; near-term option for crews to email/text photos of handwritten sheets; structured interim form if needed; later platform options include QBO Time/QuickBooks Payroll, Workyard, busybusy, ExakTime, or another field-labor system.')
        set_cell_text(row.cells[2], 'Supports Finding #8 and payroll modernization, while also improving job-costing, certified payroll data, WIP timing, labor productivity visibility, and revenue/cost cut-off.')
        break

# Add explanatory paragraph after the workflow backlog introduction.
for p in doc.paragraphs:
    if p.text.startswith('The modernization effort should document the highest-friction workflows'):
        insert_paragraph_after(p, 'Two field-data workflows should be explicit in this backlog rather than implied: daily job site reporting/photo documentation and digital timecard submission. Both are foundational controls because they convert field activity into timely evidence for project tracking, QA/QC, payroll, job costing, WIP, revenue timing, and billing readiness.')
        break

# Make Finding #17 less generic and anchor these tools there.
updates = {
    'Status: In Flight — Field Crew Communication Architecture (Quo + Notion + OpenClaw) in pilot; the four live Job Tracking databases (Installation Jobs, Inspection Jobs, Billing Log, Change Orders) provide the office-side structure.':
    'Status: In Flight — Field Crew Communication Architecture (Quo + Notion + OpenClaw) in pilot; the four live Job Tracking databases (Installation Jobs, Inspection Jobs, Billing Log, Change Orders) provide the office-side structure. Add explicit daily job site reporting/photo documentation and digital timecard submission workstreams so field communication produces usable project, payroll, job-costing, and billing evidence.',
    'Sequence: [TBD — identify predecessors / successors and whether another decision must happen first.]':
    None,
}
# Only replace the first Sequence placeholder after Finding #17.
in_f17 = False
seq_done = False
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith('Finding #17'):
        in_f17 = True
        continue
    if in_f17 and t.startswith('Finding #18'):
        in_f17 = False
    if in_f17 and t in updates and updates[t] is not None:
        p.text = updates[t]
    elif in_f17 and (not seq_done) and t == 'Sequence: [TBD — identify predecessors / successors and whether another decision must happen first.]':
        p.text = 'Sequence: Start with the lowest-friction capture method immediately (photo/email/text submissions for timecards; daily photo/report habit by foreman or crew lead), then select a durable tool once field adoption, union payroll requirements, QBO migration timing, and job-costing data needs are confirmed.'
        seq_done = True
    elif in_f17 and t == 'Success Criteria: [TBD — write as a testable done condition, not an activity.]':
        p.text = 'Success Criteria: Every active job has a daily field record containing photos or notes sufficient to understand manpower, work completed, blockers, safety issues, inspection readiness, and billing/change-order implications; every payroll period has time submitted in a format that can be tied back to employee, job, cost code/category, and payroll processing.'
    elif in_f17 and t == 'Estimated Effort: [TBD — low / moderate / high plus named resource constraints.]':
        p.text = 'Estimated Effort: Moderate. Phase 1 is process design and adoption management; Phase 2 may require vendor selection, device policy decisions, QBO/payroll integration, and field training.'
    elif in_f17 and t == 'Milestones: [TBD — dated where known; otherwise tie to transaction close, first 90 days, or Year 1.]':
        p.text = 'Milestones: 0–30 days: pilot daily reporting and photo/timecard submission on a small number of jobs. 0–90 days: decide whether CompanyCam/QBO Time/Workyard/busybusy/ExakTime or a lighter-weight workflow is the right permanent fit. Year 1: integrate with payroll, job costing, and billing/WIP reporting.'
    elif in_f17 and t == 'Open Questions: [TBD]':
        p.text = 'Open Questions: Which field employees should submit daily reports; whether company phones are required; how to handle union/certified payroll data; whether the reporting tool must integrate with QuickBooks/Notion or can operate as a standalone evidence repository.'

# Expand Finding #8 to be clearer on timecard platform options.
in_f8 = False
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith('Finding #8'):
        in_f8 = True
        continue
    if in_f8 and t.startswith('Finding #9'):
        in_f8 = False
    if in_f8 and t.startswith('Status: Scheduled — QuickBooks Desktop → QuickBooks Online with QuickBooks Payroll migration'):
        p.text = 'Status: Scheduled — QuickBooks Desktop → QuickBooks Online with QuickBooks Payroll migration, paired with a digital timecard submission path. Phase 1 may be as simple as emailed/texted photos of handwritten timesheets; later phases can evaluate QBO Time, Workyard, busybusy, ExakTime, or another field-labor system depending on union payroll, certified payroll, job-costing, and crew-adoption requirements.'
    elif in_f8 and t == 'Sequence: [TBD — identify predecessors / successors and whether another decision must happen first.]':
        p.text = 'Sequence: Do not wait for the final payroll platform to improve capture. First standardize submission timing and required fields; then align the durable tool decision with the QBO/payroll migration and job-costing design.'
    elif in_f8 and t == 'Success Criteria: [TBD — write as a testable done condition, not an activity.]':
        p.text = 'Success Criteria: Time is submitted on schedule, legible, retained, and attributable by employee/job/work type; payroll can be processed without chasing paper; certified payroll and job-costing data can be produced from the same source of truth or a controlled reconciliation.'

# Copy document properties / save.
doc.save(OUT)
print(OUT)
