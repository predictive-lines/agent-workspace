from docx import Document

SRC='efp_modernization_work/modernization-plan-efp.verify-material-check.docx'
OUT='efp_modernization_work/modernization-plan-efp.inspection-invoicing.docx'

doc=Document(SRC)
# Update workflow backlog row to explicitly call out inspection invoicing and Krissy ownership.
for table in doc.tables:
    for row in table.rows:
        if row.cells and row.cells[0].text.strip() == 'Inspection Intake → Scheduling → NFPA-25 Report → Billing / Follow-up':
            row.cells[1].text = (
                'Inbound inspection request, Krissy intake / customer coordination, Keith scheduling and inspection performance, paper NFPA-25 report preparation, return of report to the office, Krissy invoice preparation and transmittal, deficiency / follow-up handling, and customer communication.'
            )
            row.cells[2].text = (
                'Ties to inspection-delivery concentration, ITM platform decision, standing inspection agreements, and declined inspection-adjacent demand. Inspection invoicing is lower risk because Krissy already handles it, but it should still be captured as a written SOP so report completion, billing contact, invoice issuance, and follow-up do not live only in office memory.'
            )

# Add a sentence to Finding #21 because the current-state invoice delay came from missing billing data.
in21=False
for p in doc.paragraphs:
    t=p.text.strip()
    if t.startswith('Finding #21'):
        in21=True
        continue
    if in21 and t.startswith('Finding #22'):
        in21=False
    if in21 and t.startswith('Status: In Flight'):
        p.text = (
            t + ' Inspection invoicing should be included in the intake/data-model pass: even though Krissy currently handles it, the process depends on having billing contact, billing address, report status, and customer transmittal details available when the inspection is complete.'
        )
    elif in21 and t.startswith('Success Criteria:'):
        p.text = (
            'Success Criteria: Job and inspection intake records contain the minimum information required to schedule, perform, invoice, and close out the work without downstream chasing: site address, billing contact, billing address, on-site contact, report/customer transmittal details, pay terms, permit/inspection requirements where applicable, and follow-up owner.'
        )

doc.save(OUT)
print(OUT)
