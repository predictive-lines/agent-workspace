from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOCS = [
    ('efp_modernization_work/current-state-process-documentation-efp.verify-kevin.docx', 'efp_modernization_work/current-state-process-documentation-efp.styled.docx', 'Current-State Process Documentation'),
    ('efp_modernization_work/gap-analysis-efp.verify-kevin.docx', 'efp_modernization_work/gap-analysis-efp.styled.docx', 'Gap Analysis & Risk Assessment'),
    ('efp_modernization_work/modernization-plan-efp.verify-kevin.docx', 'efp_modernization_work/modernization-plan-efp.styled.docx', 'Modernization Plan & Roadmap'),
]

DARK_BLUE = RGBColor(31, 78, 121)
MID_BLUE = RGBColor(91, 155, 213)
GREY = RGBColor(89, 89, 89)
LIGHT_BLUE_FILL = 'D9EAF7'
LIGHT_GREY_FILL = 'F2F2F2'

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def set_run(run, name='Aptos', size=10, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def clear_and_set(p, text, *, align=None, size=10, bold=False, color=None, space_after=3):
    p.clear()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)

def set_style_font(style, name='Aptos', size=10, bold=None, color=None):
    font = style.font
    font.name = name
    font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if color:
        font.color.rgb = color
    style._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def style_table(table):
    # Keep existing table structure; lightly standardize grid, font, and header shading.
    try:
        table.style = 'Table Grid'
    except Exception:
        pass
    if table.rows:
        for cell in table.rows[0].cells:
            set_cell_shading(cell, LIGHT_BLUE_FILL)
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run(run, size=9, bold=True, color=DARK_BLUE)
    # Shade obvious label column in 2-col finding detail tables.
    if len(table.columns) == 2:
        for row in table.rows:
            if row.cells:
                set_cell_shading(row.cells[0], LIGHT_GREY_FILL)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    if run.font.size is None:
                        set_run(run, size=9, bold=bool(run.bold))

def standardize(src, out, title):
    doc = Document(src)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)

    styles = doc.styles
    try:
        normal_style = styles['Normal']
    except KeyError:
        normal_style = styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(normal_style, 'Aptos', 10)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.05
    for name, size, color in [
        ('Heading 1', 14, DARK_BLUE),
        ('Heading 2', 12, DARK_BLUE),
        ('Heading 3', 11, MID_BLUE),
    ]:
        if name in styles:
            set_style_font(styles[name], 'Aptos Display', size, bold=True, color=color)
            styles[name].paragraph_format.space_before = Pt(10 if name == 'Heading 1' else 6)
            styles[name].paragraph_format.space_after = Pt(4)

    # Remove leading blank paragraphs before title.
    while doc.paragraphs and not doc.paragraphs[0].text.strip():
        delete_paragraph(doc.paragraphs[0])

    # Ensure at least seven title-block paragraphs exist at top.
    while len(doc.paragraphs) < 7:
        doc.add_paragraph('')

    title_lines = [
        (title, 18, True, DARK_BLUE, WD_ALIGN_PARAGRAPH.CENTER, 1),
        ('Operational Assessment & Process Documentation', 11, False, GREY, WD_ALIGN_PARAGRAPH.CENTER, 2),
        ('Excel Fire Protection Co., Inc.', 12, True, DARK_BLUE, WD_ALIGN_PARAGRAPH.CENTER, 6),
        ('Prepared by: Predictive Lines', 10, False, None, WD_ALIGN_PARAGRAPH.CENTER, 1),
        ('Prepared for: Kevin Masich / Excel Fire Protection', 10, False, None, WD_ALIGN_PARAGRAPH.CENTER, 1),
        ('Date: May 2026', 10, False, None, WD_ALIGN_PARAGRAPH.CENTER, 1),
        ('Status: Draft for Kevin review', 10, True, None, WD_ALIGN_PARAGRAPH.CENTER, 8),
    ]
    for p, (text, size, bold, color, align, after) in zip(doc.paragraphs[:7], title_lines):
        p.style = normal_style
        clear_and_set(p, text, align=align, size=size, bold=bold, color=color, space_after=after)

    # Ensure there is visual breathing room after the title block.
    if len(doc.paragraphs) > 7 and doc.paragraphs[7].text.strip():
        spacer = doc.paragraphs[6].insert_paragraph_before('')
        # insert_paragraph_before puts it before status; remove and use raw XML after status instead
        delete_paragraph(spacer)
        new_p = OxmlElement('w:p')
        doc.paragraphs[6]._p.addnext(new_p)

    # Apply base font to all body runs without clobbering bold/italic.
    for p in doc.paragraphs[7:]:
        if p.style is None:
            p.style = normal_style
        if p.style and p.style.name in ['Heading 1', 'Heading 2', 'Heading 3']:
            for run in p.runs:
                # Heading style carries most formatting; enforce font family.
                set_run(run, name='Aptos Display', size=run.font.size.pt if run.font.size else (14 if p.style.name=='Heading 1' else 12 if p.style.name=='Heading 2' else 11), bold=bool(run.bold) or True, color=DARK_BLUE if p.style.name in ['Heading 1','Heading 2'] else MID_BLUE)
        else:
            for run in p.runs:
                if run.text:
                    set_run(run, name='Aptos', size=10, bold=bool(run.bold), color=run.font.color.rgb if run.font.color and run.font.color.rgb else None)

    for table in doc.tables:
        style_table(table)

    doc.save(out)
    print(out)

for src, out, title in DOCS:
    standardize(src, out, title)
