from docx import Document

CUR='efp_modernization_work/current-state-process-documentation-efp.kevin.docx'
MOD='efp_modernization_work/modernization-plan-efp.kevin.docx'

# Current-state: remove bracketed pending marker but keep factual note.
doc=Document(CUR)
for p in doc.paragraphs:
    if p.text.strip() == '[PENDING] Estimating rates and cost buildup methodology documentation is in progress and will be incorporated upon completion.':
        p.text = 'Estimating rates and cost-buildup methodology documentation is in progress. Kevin and Justin have planned additional working sessions to reduce the current owner/field-lead knowledge into a procedure that can be followed without Kevin or Keith’s direct assistance.'
doc.save(CUR)

# Modernization: make it read like a review draft, not an internal scaffold; replace TBD placeholders with Kevin-review language.
doc=Document(MOD)
repls={
    'Status: Draft scaffold':'Status: Draft for Kevin review',
    'This scaffold establishes the document structure, roadmap table, dependency model, and sample finding-entry format. The finding inventory is populated from the current gap-analysis draft; several workstream entries intentionally remain in template form pending Justin / Kevin / Jaclyn decisions and May on-site findings.':'This draft establishes the roadmap structure, dependency model, and finding-entry format. The finding inventory is populated from the current gap-analysis draft. Several entries intentionally remain decision-oriented pending Kevin / Justin / Jaclyn review, because assigning owners or dates without the operating decision would create false precision.',
    'Owner: [TBD]':'Owner: To confirm during Kevin / Justin review.',
    'Sequence: [TBD — identify predecessors / successors and whether another decision must happen first.]':'Sequence: To confirm during Kevin / Justin review; identify predecessors, successors, and any required operating-model decision before execution.',
    'Success Criteria: [TBD — write as a testable done condition, not an activity.]':'Success Criteria: To define as a testable done condition during Kevin / Justin review.',
    'Estimated Effort: [TBD — low / moderate / high plus named resource constraints.]':'Estimated Effort: To size after owner and workstream decision are confirmed.',
    'Milestones: [TBD — dated where known; otherwise tie to transaction close, first 90 days, or Year 1.]':'Milestones: To assign once the item is prioritized; use transaction close, first 90 days, or Year 1 gates where specific dates are not yet known.',
    'Open Questions: [TBD]':'Open Questions: Confirm owner, priority, and whether this item should be handled by process documentation, system implementation, training, or policy.',
    'Drafting note: This scaffold uses the current gap-analysis numbering as of May 12, 2026. In particular, Absence of Written QA/QC Program is Finding #20 in the live document.':'Drafting note: This draft uses the current gap-analysis numbering as of May 12, 2026. In particular, Absence of Written QA/QC Program is Finding #20 in the live document.'
}
for p in doc.paragraphs:
    if p.text in repls:
        p.text=repls[p.text]
    else:
        # defend against embedded scaffold phrasing
        if 'scaffold' in p.text.lower():
            p.text=p.text.replace('scaffold','draft').replace('Scaffold','Draft')
# Fix title metadata if in table/cells too
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if p.text in repls:
                    p.text=repls[p.text]
                elif 'scaffold' in p.text.lower():
                    p.text=p.text.replace('scaffold','draft').replace('Scaffold','Draft')

doc.save(MOD)
print('polished')
