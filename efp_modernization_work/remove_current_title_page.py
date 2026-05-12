from docx import Document
from docx.shared import Pt
from zipfile import ZipFile
from lxml import etree

SRC='efp_modernization_work/current-state-process-documentation-efp.verify-ar-cadence.docx'
OUT='efp_modernization_work/current-state-process-documentation-efp.no-title-page.docx'

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

# Load current doc.
doc=Document(SRC)

# Tighten the existing title block to match the gap-analysis first-page treatment.
# Keep content, but remove cover-page spacing and section break / blank paragraph before Section 1.
for i,p in enumerate(doc.paragraphs[:7]):
    p.paragraph_format.space_before = Pt(0)
    # compact, still readable
    if i == 0:
        p.paragraph_format.space_after = Pt(1)
    elif i == 1:
        p.paragraph_format.space_after = Pt(2)
    elif i == 2:
        p.paragraph_format.space_after = Pt(4)
    elif i < 6:
        p.paragraph_format.space_after = Pt(1)
    else:
        p.paragraph_format.space_after = Pt(6)

# Delete blank section-break paragraph immediately after the title block if present.
# This is what creates the dedicated title/cover page.
if len(doc.paragraphs) > 7 and not doc.paragraphs[7].text.strip():
    delete_paragraph(doc.paragraphs[7])

# Ensure Section 1 follows title metadata directly.
if len(doc.paragraphs) > 7 and doc.paragraphs[7].text.startswith('1. Executive Summary'):
    doc.paragraphs[7].paragraph_format.space_before = Pt(8)

# Save.
doc.save(OUT)

# Quick XML sanity check: no section break paragraph before the first heading.
with ZipFile(OUT) as z:
    root=etree.fromstring(z.read('word/document.xml'))
ns={'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
for i,pel in enumerate(root.xpath('.//w:body/w:p',namespaces=ns)[:12]):
    sect=len(pel.xpath('./w:pPr/w:sectPr',namespaces=ns))
    br=len(pel.xpath('.//w:br[@w:type="page"]',namespaces=ns))
    print(i,'sectPr',sect,'pageBreak',br)
print(OUT)
