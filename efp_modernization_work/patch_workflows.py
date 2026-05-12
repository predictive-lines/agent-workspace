from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

src=Path('efp_modernization_work/modernization-plan-efp.latest.docx')
out=Path('efp_modernization_work/modernization-plan-efp.updated.docx')
doc=Document(src)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False):
    cell.text=''
    p=cell.paragraphs[0]
    r=p.add_run(text)
    r.bold=bold
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size=Pt(8.5)

def make_table(headers, rows):
    t=doc.add_table(rows=1, cols=len(headers))
    t.style='Table Grid'
    t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], h, True)
        shade(t.rows[0].cells[i], 'D9EAF7')
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            set_cell_text(cells[i], val)
            cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
    return t

# Create subsection content at document end, then move it before Section 4.
heading=doc.add_heading('3.1 Workflow Documentation Backlog', level=2)
p=doc.add_paragraph('The modernization effort should document the highest-friction workflows both as written current-state narratives and as simple Mermaid-style flowcharts during discovery. Final versions can be rendered into graphical diagrams for Word/PDF deliverables once the flow is stable.')
rows=[
 ['Service Work Order → Pricing → Invoicing', 'Repair/service call completed; paper Service Work Order prepared by Keith or Bud; Kevin prices parts from ETNA list and calculates overhead/margin; Krissy invoices in QuickBooks.', 'Current-State §2.3; supports billing-control and service-margin modernization.'],
 ['Marketing & Sales Funnel / Inbound Opportunity Intake', 'Inbound phone calls, emails, referrals, GC/customer requests, inspection/service inquiries, and declined/deflected opportunities from first contact through routing, estimate/service decision, and whether the lead is captured anywhere.', 'Ties directly to Findings #12, #14, #15, #28, and #38; captures the current low-system, relationship-driven sales motion before CRM design.'],
 ['Installation Job Award → Contract Review → Execution → Billing/Retainage', 'LOI/contract receipt, office extraction of billing and operational requirements, scheduling, field execution, draw billing, AIA documents, lien waivers, and retainage billing.', 'Current-State §2.2–§2.3; supports PM/subcontractor-management and billing process remediation.'],
 ['Inspection Intake → Scheduling → NFPA-25 Report → Billing / Follow-up', 'Inbound inspection request, Keith scheduling, report preparation, deficiency handling, customer communication, invoice generation, and follow-up.', 'Ties to inspection-delivery concentration, ITM platform decision, standing inspection agreements, and declined inspection-adjacent demand.'],
 ['Payroll / Timecard Flow', 'Paper timesheets, field submission, AT&C payroll processing, certified payroll spreadsheet maintenance, and target QBO Payroll / digital timecard transition.', 'Supports Finding #8 and payroll modernization path.'],
 ['New Hire / Onboarding Flow', 'Local 669 dispatch, StangDS recruiting funnel, interview stages, pre-start paperwork, I-9/W-4/MI-W4, safety/training records, and onboarding checklist completion.', 'Supports Findings #2, #18, #28, #35, and #36.'],
]
table=make_table(['Workflow','Current-state flow to capture','Modernization linkage'], rows)

# Move newly appended elements before Section 4 heading.
insert_before=None
for para in doc.paragraphs:
    if para.text.strip()=='4. Resource & Effort Summary':
        insert_before=para._p
        break
if insert_before is None:
    raise SystemExit('Section 4 heading not found')
for el in [heading._p, p._p, table._tbl]:
    insert_before.addprevious(el)

doc.save(out)
print(out)
