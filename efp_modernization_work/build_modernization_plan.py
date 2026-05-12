from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
from pathlib import Path

BASE = Path(__file__).parent
GAP = BASE / 'gap-analysis-efp.docx'
OUT = BASE / 'modernization-plan-efp.docx'

STATUS_ORDER = {'Scheduled':0,'In Flight':1,'Deal Mechanics':2,'Sequenced':3,'Open':4}
HORIZON_ORDER = {'Pre-close':0,'0–90 days':1,'Year 1':2,'Strategic':3}
SEVERITY_ORDER = {'Critical':0,'High':1,'Medium':2,'Low':3}

WORKSTREAMS = {
    '3.1 Financial Operations & Controls':'Financial Operations & Controls',
    '3.2 Revenue Generation & Customer Management':'Revenue Generation & Customer Management',
    '3.3 Project Delivery & Operations':'Project Delivery & Operations',
    '3.4 People & Administration':'People & Administration',
    '3.5 Institutional Knowledge & Key Relationships':'Institutional Knowledge & Key Relationships',
    '3.6 Regulatory & Compliance':'Regulatory & Compliance',
    '3.7 Transaction-Specific Items':'Transaction-Specific Items',
}

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.size = Pt(8.5)

def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i,h in enumerate(headers):
        set_cell_text(hdr[i], h, True)
        shade(hdr[i], 'D9EAF7')
    for row in rows:
        cells = table.add_row().cells
        for i,val in enumerate(row):
            set_cell_text(cells[i], val)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    return table

def parse_gap():
    gap = Document(GAP)
    findings=[]
    current_section=None
    current_workstream=None
    n=0
    for i,p in enumerate(gap.paragraphs):
        text=p.text.strip()
        style=getattr(p.style,'name',None)
        if style=='Heading 2' and text.startswith('3.'):
            current_section=text.split(' ',1)[0]
            current_workstream=WORKSTREAMS.get(text, text)
        if style=='Heading 3' and current_workstream and text and not text.startswith(('Kevin Masich','Keith Lefebvre','Betty Laufer','Anderson')):
            # Only real findings have Severity and Remediation immediately following.
            sev_line = gap.paragraphs[i+1].text.strip() if i+1 < len(gap.paragraphs) else ''
            rem_line = gap.paragraphs[i+2].text.strip() if i+2 < len(gap.paragraphs) else ''
            if not sev_line.startswith('Severity:') or not rem_line.startswith('Remediation:'):
                continue
            n+=1
            m = re.match(r'Severity:\s*([^│]+)│\s*Type:\s*([^│]+)│\s*Horizon:\s*([^│]+)│\s*Current-State:\s*(.+)$', sev_line)
            severity=typ=horizon=csp=''
            if m:
                severity=m.group(1).strip(); typ=m.group(2).strip(); horizon=m.group(3).strip(); csp=m.group(4).strip()
            rm = re.match(r'Remediation:\s*([^—]+)\s*—\s*(.*)$', rem_line)
            status = rm.group(1).strip() if rm else ''
            remediation = rm.group(2).strip() if rm else rem_line.replace('Remediation:','').strip()
            findings.append({
                'num': n,
                'title': text,
                'section': current_section,
                'gap_ref': f'Gap Analysis §{current_section}',
                'workstream': current_workstream,
                'severity': severity,
                'type': typ,
                'horizon': horizon,
                'csp': csp,
                'status': status,
                'remediation': remediation,
            })
    return findings

findings=parse_gap()
by_num={f['num']:f for f in findings}

doc=Document()
sec=doc.sections[0]
sec.top_margin=Inches(0.7); sec.bottom_margin=Inches(0.7); sec.left_margin=Inches(0.7); sec.right_margin=Inches(0.7)
styles=doc.styles
styles['Normal'].font.name='Aptos'; styles['Normal'].font.size=Pt(10)
for s in ['Heading 1','Heading 2','Heading 3']:
    styles[s].font.name='Aptos'

title=doc.add_paragraph()
title.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=title.add_run('Modernization Plan & Roadmap')
r.bold=True; r.font.size=Pt(18)
for line in ['Operational Assessment & Process Documentation','Excel Fire Protection Co., Inc.','Prepared by: Predictive Lines','Prepared for: Kevin Masich / Excel Fire Protection','Date: May 2026','Status: Draft scaffold']:
    p=doc.add_paragraph(line); p.alignment=WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_heading('1. Purpose & Scope', level=1)
doc.add_paragraph('This document is the Modernization Plan & Roadmap deliverable contemplated by the March 30, 2026 Statement of Work. It is designed to read alongside, not replace, the Current-State Process Documentation and Gap Analysis & Risk Assessment.')
doc.add_paragraph('The gap analysis remains the dated diagnostic snapshot. This roadmap is the action document: it translates findings into sequencing, ownership, success criteria, dependencies, effort, and milestones. As remediation progresses, this document should move; the diagnostic text in the gap analysis should remain comparatively stable.')
doc.add_paragraph('This scaffold establishes the document structure, roadmap table, dependency model, and sample finding-entry format. The finding inventory is populated from the current gap-analysis draft; several workstream entries intentionally remain in template form pending Justin / Kevin / Jaclyn decisions and May 2026 on-site outcomes.')

doc.add_heading('2. Roadmap Overview', level=1)
doc.add_paragraph('The table below groups the full finding inventory by the remediation horizon assigned in the gap analysis. Within each horizon, items are sorted by severity and current remediation status so the work can be reviewed as an execution queue rather than as a narrative report.')
rows=[]
for f in sorted(findings, key=lambda x:(HORIZON_ORDER.get(x['horizon'],99), SEVERITY_ORDER.get(x['severity'],99), STATUS_ORDER.get(x['status'],99), x['num'])):
    rows.append([f"#{f['num']}", f['horizon'], f['severity'], f['status'], f['title'], f['workstream']])
add_table(doc, ['#','Horizon','Severity','Status','Finding','Workstream'], rows)

doc.add_paragraph('Narrative summary. Pre-close work is dominated by visible-state compliance and transaction-mechanics items: new-hire form cleanup, UIA reconstruction, office succession, BFS certification verification, and the sunset of the legacy banking signature-stamp practice. The first 90 days focus on business-continuity risks: inspection-delivery concentration, field-to-office communication, IT backup/security, safety program management, and compliance tracking. Year 1 shifts toward structural modernization: QBO/payroll, estimating documentation, inspection-service strategy, QA/QC, project-management discipline, customer/revenue segmentation, MDM, fleet and inventory hygiene, and management reporting. Strategic items should not be forced into the first 90 days unless they become prerequisites for revenue protection or buyer/lender requirements.')

doc.add_heading('3. Sequencing & Dependencies', level=1)
doc.add_paragraph('The roadmap has several genuine dependencies. These should be managed explicitly; otherwise the team will create documents or tools before the operating model underneath them is settled.')
rows=[
 ['Inspection-delivery hire decision', 'Unlocks Finding #16 directly and materially affects Findings #12, #15, #23, #24, #38, and #40. Decide whether the mitigation path is Konner promotion/cross-training, external NFPA-25 hire, or a hybrid.'],
 ['Deal close / buy-in mechanics', 'Unlocks Findings #1 and #3 through new banking relationships, authority tiers, and dual-control thresholds in the shareholder or operating agreement.'],
 ['QuickBooks Desktop → QBO / Payroll migration', 'Provides the timing window to restructure expense categories, revenue segmentation, certified-payroll tooling, customer master data, and recurring management reporting without rebuilding twice.'],
 ['HR / compliance calendar buildout', 'Common backbone for safety program management, ISN obligations, customer-specific compliance tracking, COI expirations, certification/license tracking, and training records.'],
 ['Field Crew Communication + Job Tracking adoption', 'Provides the operating substrate for job intake, rough-in/final readiness, document centralization, permit workflow, and field-to-office communication discipline.'],
 ['Estimating methodology capture', 'Highest-leverage owner-knowledge dependency; also supports sales training, pricing discipline, budget forecasting, and eventual delegation of estimate preparation.'],
]
add_table(doc, ['Dependency / Critical Path Item','Downstream effect'], rows)

doc.add_heading('4. Resource & Effort Summary', level=1)
doc.add_paragraph('This section should become a lightweight capacity model, not a consulting-theater estimate. The immediate practical question is whether the remediation queue can fit inside Justin, Jaclyn, Kevin, Krissy, AT&C, Lee Init, Gauthier, StangDS, and field-leadership capacity without starving day-to-day operations.')
rows=[
 ['Justin / Predictive Lines', 'Roadmap owner; document/process design; Notion + OpenClaw implementation; operating-model decisions; May 11/18 working sessions.', 'High through first 90 days, tapering in Year 1.'],
 ['Kevin', 'Source of owner knowledge, estimating methodology, customer relationships, banking authority transition, PM/subcontractor-management discussion.', 'High but must be tightly scheduled; use interviews/workshops, not open-ended asks.'],
 ['Jaclyn', 'Finance/control environment, QBO/payroll migration decisions, reporting cadence, approval authority, vendor/payment controls.', 'Moderate; spikes around finance migration and internal controls.'],
 ['AT&C', 'UIA reconstruction, payroll/tax support, QuickBooks transition support, chart-of-accounts cleanup.', 'Moderate; external dependency and possible bottleneck.'],
 ['External vendors', 'Lee Init / IT, Gauthier Insurance, StangDS, potential ITM platform vendor, possible HR/payroll support.', 'Use where specialized capacity prevents Justin/Kevin bottlenecks.'],
]
add_table(doc, ['Resource','Likely role','Capacity note'], rows)

doc.add_heading('5. Workstream-Aligned Remediation Plans', level=1)
doc.add_paragraph('Each entry below should stay short and action-shaped. The format deliberately avoids re-stating the diagnostic prose from the gap analysis. The diagnosis lives there; this section should answer: who owns it, what depends on it, how we know it is done, and what happens next.')

sample_details={
2:{
'owner':'Justin / Krissy, with Kevin support as needed',
'sequence':'Pre-close scheduled item; feeds the standing new-hire onboarding process and employee-records cleanup.',
'success':'For every current employee, either a complete I-9, W-4, and MI-W4 is located and filed, or replacement forms are collected and retained in the employee record. The New Employee Onboarding Checklist becomes the forward process so the gap does not recur.',
'effort':'Low to moderate; one focused office audit plus employee follow-up for missing forms.',
'milestones':'May 11 on-site: audit existing files and identify gaps. Week of May 11: collect replacements. Post-close: spot-check onboarding checklist against next new hire.',
'questions':'Confirm where employee files will live long-term: Paperless / controlled Drive / M365 / other canonical repository.'},
11:{
'owner':'Justin + Kevin, with Keith input and Scott benchmarking input if the Wisconsin visit occurs.',
'sequence':'Highest-leverage owner-knowledge capture item. Should happen before broader sales process or estimating delegation decisions.',
'success':'A documented estimating process exists that can be followed without Kevin or Keith’s assistance, including cost buildup, labor assumptions, materials/subcontractor inputs, markup/contingency logic, review/approval points, and examples from recent bids.',
'effort':'Moderate to high; two working sessions plus one test pass by someone other than Kevin/Keith.',
'milestones':'Week of May 11: map current method and collect example estimates. Week of May 18: draft v0 and test against one historical job. Follow-up: revise based on test failure points.',
'questions':'Decide whether the first version should be a Word/Notion SOP only, or include spreadsheet templates / calculators as controlled artifacts.'},
20:{
'owner':'Not yet assigned; likely Justin + Keith/Kevin initially, then inspection/project delivery owner once named.',
'sequence':'Open item. Should not be overbuilt until the inspection-delivery and PM ownership decisions are clearer, but the absence of a QA/QC baseline is a High finding and needs an owner.',
'success':'A written QA/QC program defines minimum review points, inspection/readiness checks, documentation artifacts, correction tracking, and accountability for install and inspection work. The program is light enough to be used by the field, not just filed.',
'effort':'Moderate; best built from actual rework/callback patterns and AHJ re-trip causes rather than from a generic template.',
'milestones':'May on-site: identify current informal QA checks and common failure modes. 0–90 days: name owner and define minimum QA gates. Year 1: fold QA artifacts into job tracking / permit / inspection workflows.',
'questions':'Who owns QA/QC after Keith’s eventual transition: field lead, inspector, office PM function, or new operations manager?'}
}

for ws in [v for k,v in WORKSTREAMS.items() if v != 'Transaction-Specific Items']:
    doc.add_heading(f"5.{list(WORKSTREAMS.values()).index(ws)+1} {ws}", level=2)
    for f in [x for x in findings if x['workstream']==ws]:
        doc.add_heading(f"Finding #{f['num']} — {f['title']} ({f['severity']})", level=3)
        lines=[
            ('Cross-reference', f"{f['gap_ref']}; Current-State {f['csp']}"),
            ('Status', f"{f['status']} — {f['remediation']}"),
        ]
        detail=sample_details.get(f['num'])
        if detail:
            lines += [('Owner',detail['owner']),('Sequence',detail['sequence']),('Success Criteria',detail['success']),('Estimated Effort',detail['effort']),('Milestones',detail['milestones']),('Open Questions',detail['questions'])]
        else:
            lines += [('Owner','[TBD]'),('Sequence','[TBD — identify predecessors / successors and whether another decision must happen first.]'),('Success Criteria','[TBD — write as a testable done condition, not an activity.]'),('Estimated Effort','[TBD — low / moderate / high plus named resource constraints.]'),('Milestones','[TBD — dated where known; otherwise tie to transaction close, first 90 days, or Year 1.]'),('Open Questions','[TBD]')]
        for label,value in lines:
            p=doc.add_paragraph()
            r=p.add_run(label+': '); r.bold=True
            p.add_run(value)

doc.add_heading('6. Items Not Yet Scoped', level=1)
doc.add_paragraph('The following open items do not yet have a sufficiently concrete modernization workstream. Each needs a decision trigger before the plan should pretend there is an execution path.')
open_rows=[]
for f in findings:
    if f['status']=='Open':
        trigger='Assign an owner and decide whether to fold into an existing workstream or defer intentionally.'
        if f['num']==20: trigger='Decide QA/QC ownership after the field-leadership / inspection-delivery discussion.'
        elif f['num'] in (41,42): trigger='Confirm HR/compliance calendar scope and owner.'
        elif f['num'] in (4,5,6,10): trigger='Tie to QBO migration / finance modernization scope.'
        elif f['num'] in (13,15): trigger='Decide inspection-service strategy and whether to measure declined inbound volume first.'
        open_rows.append([f"#{f['num']}", f['title'], f['horizon'], trigger])
add_table(doc, ['#','Finding','Horizon','Decision trigger'], open_rows)

doc.add_heading('7. Out-of-Scope Items', level=1)
doc.add_paragraph('The modernization plan should not become a catch-all business plan. The items below may matter strategically, but they should be tracked separately unless they directly remediate a gap-analysis finding.')
for item in ['New product-line investigations not tied to an identified gap.', 'Marquette warehouse / facility expansion options except where required for document retention, backup, or operations continuity.', 'Sister-company concepts such as extinguisher service unless Justin chooses to tie them to declined inspection-adjacent inbound demand.', 'Long-range M&A / buyer-readiness positioning beyond maintaining the gap analysis as a dated diagnostic artifact.']:
    doc.add_paragraph(item, style='List Bullet')

# Footer note-ish final paragraph
p=doc.add_paragraph()
r=p.add_run('Drafting note: '); r.bold=True
p.add_run('This scaffold uses the current gap-analysis numbering as of May 12, 2026. In particular, Absence of Written QA/QC Program is Finding #20 in the live document.')

doc.save(OUT)
print(f'Wrote {OUT} with {len(findings)} findings')
