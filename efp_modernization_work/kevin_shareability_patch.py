from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

CUR_SRC='efp_modernization_work/current-state-process-documentation-efp.latest.docx'
GAP_SRC='efp_modernization_work/gap-analysis-efp.latest.docx'
MOD_SRC='efp_modernization_work/modernization-plan-efp.latest2.docx'
CUR_OUT='efp_modernization_work/current-state-process-documentation-efp.kevin.docx'
GAP_OUT='efp_modernization_work/gap-analysis-efp.kevin.docx'
MOD_OUT='efp_modernization_work/modernization-plan-efp.kevin.docx'

def insert_paragraph_after(paragraph, text, style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    p = paragraph._parent.add_paragraph()
    p._p = new_p
    p._element = new_p
    if style:
        p.style = style
    p.add_run(text)
    return p

def set_cell_text(cell, text):
    cell.text=''
    p=cell.paragraphs[0]
    r=p.add_run(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size=Pt(9)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP

def patch_current():
    doc=Document(CUR_SRC)
    for p in doc.paragraphs:
        t=p.text.strip()
        if t.startswith('Employee records are sparse and are maintained in paper form.'):
            p.text=(
                'Employee records are sparse and are maintained in paper form. The contents of the personnel files, where they exist, are not consistent across employees: the predominant content is physical timesheets retained from prior pay periods, with other materials present on a per-employee basis without a defined file standard. Required federal and state new-hire documentation has not yet been confirmed on an employee-by-employee basis; the owner has flagged that I-9 employment eligibility verification forms, W-4 federal withholding certificates, and Michigan W-4 (MI-W4) state withholding certificates may not be uniformly present and plans to verify the personnel files directly. Beyond the new-hire paperwork, the Company does not systematically retain copies of training completions, journeyman or apprentice certifications, Local 669 membership cards, mine-site safety training certificates, OSHA 10/30 cards, or NFPA continuing-education records; where such artifacts exist they live with the individual employee or in the Local 669 / NASI administrative ecosystem rather than in a Company-controlled employee file.'
            )
        elif t == '[PENDING] Verification of BFS (Bureau of Fire Services) Act 144 certification status is pending.':
            p.text=(
                'Excel Fire Protection Co., Inc. holds Bureau of Fire Services fire suppression firm certificate S-0440, valid through September 2028. The certificate authorizes the Company to design, maintain, service, or document the installation and modification of required fire suppression systems in the following categories: J — Pre-engineered foam systems; K — Custom designed, water-based sprinkler systems, including underground water supply; and P — Engineered pump pressure supply. The certificate is issued under the Public Act 207 / BFS firm-certification framework historically referred to as Act 144 certification. Kevin Masich is the current qualifying person associated with the certificate. The current-state continuity issue is therefore not absence of firm certification, but the concentration of qualifying-person status and related licensing knowledge in Kevin.'
            )
        elif t.startswith('Kevin Masich is the sole holder of the Michigan mechanical contractor license'):
            p.text=(
                'Kevin Masich is the sole holder of the Michigan mechanical contractor license with a fire suppression specialty. This license is required for the Company to operate as a fire sprinkler contractor in the state. No other employee holds this license. Kevin is also the qualifying person tied to Excel Fire Protection Co., Inc.’s BFS fire suppression firm certificate S-0440, which creates a related regulatory continuity dependency if Kevin is no longer available to serve in that capacity.'
            )
        elif t.startswith('The office maintains a physical whiteboard listing active jobs'):
            p.text=(
                'The office maintains a physical whiteboard listing active jobs, the dollar amount outstanding, and a percentage-complete estimate. The whiteboard is updated informally by Betty. It is the Company’s primary visual artifact for “what is in progress right now and how far along is it,” and substitutes for a project management or work-in-process tracking system. As noted in the field-to-office status cadence discussion above and in Section 2.1, the whiteboard’s accuracy is directly bounded by what the field has reported back to the office; in practice it is frequently out of date in both directions (jobs that have advanced without notification sit too low, and jobs that have been completed without notification sit unresolved), and it should be treated as a directional reference rather than a system of record. The Company does not currently use a daily job-site reporting tool, photo documentation system, or structured daily foreman report process comparable to CompanyCam; field progress, site photos, blockers, safety observations, inspection-readiness status, and change-order signals are not captured in a consistent daily artifact.'
            )
    doc.save(CUR_OUT)

def patch_gap():
    doc=Document(GAP_SRC)
    for p in doc.paragraphs:
        t=p.text.strip()
        if t.startswith('Remediation: Scheduled — Workforce-wide audit'):
            p.text='Remediation: Scheduled — Workforce-wide personnel-file audit and replacement-form collection; Justin plans to verify the personnel files directly on May 12. The New Employee Onboarding Checklist (Notion) is the standing process for new hires going forward.'
        elif t.startswith('The owner reports that Form I-9 employment eligibility'):
            p.text=(
                'The owner has flagged that Form I-9 employment eligibility verification documents, IRS Form W-4 federal withholding certificates, and Michigan W-4 (MI-W4) state withholding certificates may not be on file for every current employee. Personnel files, where they exist, contain principally physical timesheets retained from prior pay periods, with new-hire documentation present on a per-employee basis without a defined file standard. This finding should be read as a known verification risk until the personnel-file audit confirms which forms are present and which, if any, require replacement.'
            )
        elif t.startswith('Pre-close remediation should include a workforce-wide audit'):
            p.text=(
                'Pre-close remediation should include a workforce-wide audit of personnel files, completion of replacement I-9 forms where missing (the Form I-9 instructions provide for a corrective procedure where the original form has been lost or never completed), and replacement W-4 / MI-W4 collections where missing. If the May 12 file review confirms that some or all forms are already present, this finding should be narrowed to the actual missing-document population rather than left as a blanket deficiency.'
            )
        elif 'QA/QC (Finding #14, open)' in t:
            p.text=t.replace('QA/QC (Finding #14, open)', 'QA/QC (Finding #20, open)')
        elif t.startswith('BFS Act 144 Certification Verification'):
            p.text='BFS / S-0440 Fire Suppression Firm Certification Continuity'
        elif t.startswith('Remediation: Scheduled — On the May 11 on-site information-gathering list.'):
            p.text='Remediation: Verified / Sequenced — Certificate confirmed onsite; continuity planning remains because Kevin is the current qualifying person.'
        elif t.startswith('Verification of Bureau of Fire Services Act 144 certification status remains pending'):
            p.text=(
                'Bureau of Fire Services fire suppression firm certification has been confirmed. Excel Fire Protection Co., Inc. holds certificate S-0440, valid through September 2028, authorizing design, maintenance, service, or documentation of the installation and modification of required fire suppression systems in categories J (pre-engineered foam systems), K (custom designed, water-based sprinkler systems, including underground water supply), and P (engineered pump pressure supply). The residual risk is succession and continuity: Kevin Masich is the current qualifying person, and no backup qualifying person or owner licensing/certification path is yet documented.'
            )
        elif t.startswith('BFS Act 144 certification verification. Confirm and integrate'):
            p.text='BFS / S-0440 certification continuity. Certificate has been verified; document the qualifying-person succession path and owner licensing/certification plan.'
        elif t.startswith('Sole holder of the Michigan mechanical contractor license'):
            p.text=(
                'Sole holder of the Michigan mechanical contractor license with fire suppression specialty and current qualifying person for Excel Fire Protection Co., Inc.’s BFS fire suppression firm certificate S-0440. Sole holder of estimating methodology and cost-buildup rates. Primary or sole point of contact for most of the Company’s significant customer relationships. Primary decision authority on pricing, scope, and customer commitments. Sole signature on the bank’s signature card for the operating account. Financial oversight is exercised personally rather than through a reporting structure.'
            )
    doc.save(GAP_OUT)

def patch_mod():
    doc=Document(MOD_SRC)
    for p in doc.paragraphs:
        t=p.text.strip()
        if t.startswith('Narrative summary. Pre-close work is dominated'):
            p.text=t.replace('BFS certification verification', 'BFS/S-0440 certification continuity').replace('new-hire form cleanup', 'new-hire form verification/cleanup')
        elif 'QA/QC (Finding #14, open)' in t:
            p.text=t.replace('QA/QC (Finding #14, open)', 'QA/QC (Finding #20, open)')
        elif t.startswith('Finding #44 — BFS Act 144 Certification Verification'):
            p.text='Finding #44 — BFS / S-0440 Fire Suppression Firm Certification Continuity (Low)'
        elif t.startswith('Status: Scheduled — On the May 11 on-site information-gathering list.'):
            p.text='Status: Verified / Sequenced — Excel Fire Protection Co., Inc. certificate S-0440 confirmed valid through September 2028 for categories J/K/P; continuity work remains because Kevin is the current qualifying person.'
        elif t == 'Owner: [TBD]' and False:
            pass
    # Update #44 block placeholders only
    in44=False
    for p in doc.paragraphs:
        t=p.text.strip()
        if t.startswith('Finding #44'):
            in44=True; continue
        if in44 and t.startswith('6. Items Not Yet Scoped'):
            in44=False
        if in44:
            if t == 'Owner: [TBD]': p.text='Owner: Justin, with Kevin input; BFS/LARA as external regulatory source.'
            elif t.startswith('Sequence: [TBD'): p.text='Sequence: Certificate verification is complete. Next step is to map Kevin’s qualifying-person role, confirm BFS notice/replacement requirements, and build an owner/additional-employee licensing and certification path so S-0440 continuity is not dependent on one person.'
            elif t.startswith('Success Criteria: [TBD'): p.text='Success Criteria: Excel Fire maintains S-0440 certification without interruption; at least one backup path is documented for mechanical contractor fire suppression licensing and BFS qualifying-person coverage; expiration, renewal, and notice deadlines are tracked in the compliance calendar.'
            elif t.startswith('Estimated Effort: [TBD'): p.text='Estimated Effort: Moderate; mostly requirements research, LARA/BFS correspondence, application/exam planning, and internal candidate development.'
            elif t.startswith('Milestones: [TBD'): p.text='Milestones: May 2026: capture certificate and Kevin qualifying-person facts. 0–90 days: confirm Justin licensing/experience-waiver path and BFS qualifying-person requirements. Year 1: begin execution of owner/backup certification plan.'
            elif t == 'Open Questions: [TBD]': p.text='Open Questions: Which associate degree programs, if any, qualify for the one-year Michigan mechanical contractor experience waiver; what experience/training BFS will accept for S-0440 qualifying-person status; which internal employee could become the second qualified option.'
    # Update table entries in roadmap table.
    for table in doc.tables:
        for row in table.rows:
            vals=[c.text.strip() for c in row.cells]
            if vals and vals[0]=='#44':
                if len(vals)>3: set_cell_text(row.cells[3], 'Verified / Sequenced')
                if len(vals)>4: set_cell_text(row.cells[4], 'BFS / S-0440 Fire Suppression Firm Certification Continuity')
            for cell in row.cells:
                if 'BFS Act 144 Certification Verification' in cell.text:
                    set_cell_text(cell, cell.text.replace('BFS Act 144 Certification Verification','BFS / S-0440 Fire Suppression Firm Certification Continuity'))
                if 'Finding #14, open' in cell.text:
                    set_cell_text(cell, cell.text.replace('Finding #14, open','Finding #20, open'))
    doc.save(MOD_OUT)

patch_current(); patch_gap(); patch_mod()
print(CUR_OUT); print(GAP_OUT); print(MOD_OUT)
