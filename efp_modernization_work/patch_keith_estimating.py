from docx import Document

CUR='efp_modernization_work/current-state-process-documentation-efp.verify-styled.docx'
GAP='efp_modernization_work/gap-analysis-efp.verify-styled.docx'
MOD='efp_modernization_work/modernization-plan-efp.verify-styled.docx'
CUR_OUT='efp_modernization_work/current-state-process-documentation-efp.estimating.docx'
GAP_OUT='efp_modernization_work/gap-analysis-efp.estimating.docx'
MOD_OUT='efp_modernization_work/modernization-plan-efp.estimating.docx'

# Current-state: enrich as-is estimating methodology with Keith notes.
doc=Document(CUR)
for p in doc.paragraphs:
    t=p.text.strip()
    if t.startswith('Both Kevin (owner) and Keith (journeyman/field lead) are capable of preparing estimates'):
        p.text=(
            'Both Kevin (owner) and Keith (journeyman/field lead) are capable of preparing estimates and bids, though Keith now appears to carry much of the working estimating method for sprinkler layout, material quantity, and labor production. Two primary estimating methods are used in parallel on every job, and the two methods are intentionally cross-checked against one another as a sanity check: a dollars-per-square-foot method based on the building footprint, and a dollars-per-sprinkler-head method based on the projected head count. Both numbers are worked up for the same project and compared; a meaningful discrepancy between the two triggers a closer review of the project plans.'
        )
    elif t.startswith('The review process specifically looks for features'):
        p.text=(
            'The review process specifically looks for features that would drive the head count or labor off the per-square-foot benchmark: wall coring, ceiling obstructions, unusual room geometry, high ceilings, concealed spaces, and similar conditions. Keith’s working approach begins by estimating the number of sprinkler heads required based on the space, layout, and visible constraints in the plans. Materials are then approximated from the head count and pipe size requirements; Keith uses a rule of thumb of approximately 15 feet of pipe per head, adjusted based on the pipe size and layout conditions. Labor is estimated from field production experience, using approximately five heads installed per person per day as the core productivity assumption, then converting the head count into crew-days. The final bid figure is reconciled on a sheet that the owner refers to as the “bid sheet” — a working document used during bid prep to lay out initial square footage, head count, both unit prices, and the final total. The bid sheet is the artifact that captures the cross-check; it is prepared by hand at the drafting table alongside the plans. The process is disciplined in practice but remains experience-based rather than formalized in a standalone estimating standard, calculator, or written procedure. Kevin’s role, based on the most recent field interview notes, is more review-oriented than hands-on for this method: he reviews the estimate rather than independently rebuilding the head-count, pipe, and labor assumptions from scratch.'
        )
doc.save(CUR_OUT)

# Gap: sharpen the finding around the specific undocumented heuristics.
doc=Document(GAP)
for p in doc.paragraphs:
    t=p.text.strip()
    if t.startswith('The Company’s two-method cross-check estimating approach'):
        p.text=(
            'The Company’s two-method cross-check estimating approach (dollars-per-square-foot and dollars-per-sprinkler-head, reconciled on a working bid sheet) is consistent and appears to be disciplined, but the underlying rates, cost buildup methodology, and review heuristics are held in Keith’s and Kevin’s working knowledge rather than in written documentation. Keith’s current method includes estimating the sprinkler head count from the space and layout, approximating material from the head count and pipe size using roughly 15 feet of pipe per head as a rule of thumb, and estimating labor by converting head count into crew-days using approximately five heads installed per person per day. Kevin appears to function primarily as a reviewer of the resulting estimate rather than as the main driver of the detailed head-count, material, and labor buildup. This is a key-person dependency as much as it is a process gap, and is the single highest-leverage knowledge-transfer item in the engagement: estimating capability is what allows the Company to bid work, and an unbacked-up estimator is a single-event business-interruption risk under the closing timeline.'
        )
doc.save(GAP_OUT)

# Modernization: add the heuristics to Finding #11 success criteria / milestones.
doc=Document(MOD)
in11=False
for p in doc.paragraphs:
    t=p.text.strip()
    if t.startswith('Finding #11'):
        in11=True
        continue
    if in11 and t.startswith('Finding #12'):
        in11=False
    if in11 and t.startswith('Status: Scheduled'):
        p.text=(
            'Status: Scheduled — Focus weeks of May 11 and May 18; success criterion is a documented process that can be followed without Kevin or Keith’s assistance. The draft should capture Keith’s current head-count, pipe-per-head, and heads-per-day labor heuristics and benchmark them against Scott’s more modern shop if the Wisconsin visit occurs.'
        )
    elif in11 and t.startswith('Success Criteria: A documented estimating process exists'):
        p.text=(
            'Success Criteria: A documented estimating process exists that can be followed without Kevin or Keith’s assistance, including: plan review inputs; sprinkler head count from space/layout; material buildup by head count, pipe size, and Keith’s approximate 15-feet-of-pipe-per-head rule of thumb; labor buildup using approximately five heads installed per person per day; subcontractor/material inputs; markup/contingency logic; Kevin review/approval points; and examples from recent bids.'
        )
    elif in11 and t.startswith('Milestones: Week of May 11'):
        p.text=(
            'Milestones: Week of May 11: map Keith’s current method and collect example estimates, including head-count assumptions, pipe-size/material assumptions, 15-feet-per-head rule of thumb, and five-heads-per-person-per-day labor assumption. Week of May 18: draft v0 and test against one historical job. Follow-up: revise based on test failure points and decide whether to build a controlled estimating calculator/template.'
        )
    elif in11 and t.startswith('Open Questions: Decide whether'):
        p.text=(
            'Open Questions: Decide whether the first version should be a Word/Notion SOP only, or include spreadsheet templates / calculators as controlled artifacts. Confirm when Kevin materially adjusts Keith’s estimate versus simply reviewing it, and document those review thresholds.'
        )
doc.save(MOD_OUT)
print(CUR_OUT); print(GAP_OUT); print(MOD_OUT)
