from docx import Document
SRC='efp_modernization_work/modernization-plan-efp.verify-qbo-payroll-transition.docx'
OUT='efp_modernization_work/modernization-plan-efp.after-close-safe.docx'
doc=Document(SRC)

def patch_section(heading, replacements):
    start=None; end=len(doc.paragraphs)
    for i,p in enumerate(doc.paragraphs):
        if p.text.strip()==heading:
            start=i
            break
    if start is None:
        raise SystemExit(f'Heading not found: {heading}')
    for j in range(start+1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name.startswith('Heading 3') and doc.paragraphs[j].text.strip().startswith('Finding #'):
            end=j; break
    for label,new_text in replacements.items():
        hit=False
        for p in doc.paragraphs[start:end]:
            if p.text.strip().startswith(label):
                p.text=new_text
                hit=True
                break
        if not hit:
            raise SystemExit(f'Label {label} not found in {heading}')

patch_section('Finding #1 — Banking Signature-Stamp Practice (Critical)', {
 'Sequence:': 'Sequence: Treat as a close / immediate post-close control conversion. Confirm the closing-bank operating model, replace the legacy signature-stamp practice with named signers and documented authority thresholds, update online-banking administrator access, issue or cancel cards/check stock as needed, and align signer changes with the shareholders’/operating agreement authority tiers. This should be coordinated with the signing-to-close transition checklist, but durable spending authority and banking controls belong in the modernization plan after close.',
 'Success Criteria:': 'Success Criteria: Kevin’s signature-stamp practice is no longer used; bank signers and online-banking administrators match the post-close authority model; cardholders, limits, check controls, and approval thresholds are documented; and Betty/current-office access is either removed, retained only for an approved transition role, or converted to named-role access with appropriate controls.',
 'Milestones:': 'Milestones: Close week: confirm bank transition appointment and required signer documentation. Day 1 / first 30 days: update signers, online banking, cardholders, and approval limits. 0–90 days: test dual-control workflow for vendor payments, payroll funding, and unusual disbursements; reconcile controls against the shareholders’/operating agreement.',
 'Open Questions:': 'Open Questions: Which bank account structure will be used at close; whether old accounts are retained or replaced; who will have online-banking admin rights; what dollar thresholds require Justin/Jaclyn/Kevin approval; how long, if at all, Betty retains any access during transition; and which controls are bank-configured versus internal policy.'
})

patch_section('Finding #30 — IT Security Posture Beyond the Email Account (High)', {
 'Sequence:': 'Sequence: After close, move 2FA and account-recovery paths off Betty-personal channels and any other individual-only recovery paths. Establish named administrator accounts, shared/recoverable business-controlled recovery methods, and password-manager ownership before broader system cleanup. Coordinate with AT&C for M365, QuickBooks/Desktop infrastructure, endpoint access, and any vendor portals tied to office operations.',
 'Success Criteria:': 'Success Criteria: No critical system depends on Betty’s personal phone/email or an undocumented individual recovery path; M365 and other core systems have at least two named business-controlled administrators; 2FA methods are documented and recoverable; passwords/secrets are held in the approved password manager rather than personal notes or browser storage; and access removal/retention decisions are tied to the post-close transition role.',
 'Milestones:': 'Milestones: Close / first 30 days: inventory current 2FA and recovery methods for M365, QuickBooks/desktop access, payroll/tax portals, banking, insurance, vendor portals, and utilities. 0–90 days: move recovery methods to business-controlled channels, add backup admins, and test account recovery. Year 1: fold this into recurring access review and offboarding controls.',
 'Open Questions:': 'Open Questions: Which systems currently use Betty’s phone/email for 2FA or recovery; what business-controlled phone/email should receive recovery traffic; which accounts require AT&C intervention; which access should remain during any Betty transition role; and how often admin/recovery access should be reviewed.'
})

doc.save(OUT)
print(OUT)
