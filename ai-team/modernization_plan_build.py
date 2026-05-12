from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = 'ai-team/modernization-plan-efp.docx'

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

def add_meta_line(doc, label, value):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    p.add_run(value)


def add_status_table(doc, rows):
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    headers = ['Finding', 'Horizon', 'Status', 'Owner', 'Dependency', 'Modernization Plan Treatment']
    for c, h in zip(hdr, headers):
        set_cell_text(c, h, True)
        set_cell_shading(c, 'D9EAF7')
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell_text(cell, value)
    doc.add_paragraph()


def add_entry(doc, num, title, severity, gap_ref, status, owner, sequence, criteria, effort, milestones, questions):
    doc.add_heading(f'Finding #{num} — {title} ({severity})', level=3)
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    fields = [
        ('Gap Analysis Cross-Reference', gap_ref),
        ('Status', status),
        ('Owner', owner),
        ('Sequence', sequence),
        ('Success Criteria', criteria),
        ('Estimated Effort', effort),
        ('Milestones', milestones),
        ('Open Questions / Decisions', questions),
    ]
    for label, value in fields:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, True)
        set_cell_shading(cells[0], 'F2F2F2')
        set_cell_text(cells[1], value)
    doc.add_paragraph()


def add_template_entry(doc, num, title, severity, gap_ref='Gap Analysis §[x.x]', horizon='[Pre-close / 0–90 days / Year 1 / Strategic]', status='[Scheduled / In Flight / Sequenced / Deal Mechanics / Open]', owner='[Justin / Jaclyn / Kevin / AT&C / external / TBD]'):
    add_entry(
        doc, num, title, severity, gap_ref, status, owner,
        f'Horizon: {horizon}. Predecessors: [TBD]. Successors / unlocked findings: [TBD].',
        '[Write a testable end state. Example: “Process can be performed by a non-owner using the documented procedure and retained evidence.”]',
        '[T-shirt size or rough hours; identify whether effort is consultant, company, AT&C, or vendor time.]',
        '[Dates or gates. Use “decision trigger” if no date exists yet.]',
        '[Open questions, owner assignment, cost approval, policy decision, or source document needed.]'
    )


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.7)
sec.right_margin = Inches(0.7)

styles = doc.styles
styles['Normal'].font.name = 'Aptos'
styles['Normal'].font.size = Pt(10)
for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
    styles[style_name].font.name = 'Aptos Display'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Excel Fire Protection')
r.bold = True
r.font.size = Pt(18)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Modernization Plan & Roadmap')
r.bold = True
r.font.size = Pt(16)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Draft scaffold for Predictive Lines consulting engagement')

add_meta_line(doc, 'Prepared for: ', 'Excel Fire Protection / Kevin Masich')
add_meta_line(doc, 'Prepared by: ', 'Predictive Lines / Justin Miller')
add_meta_line(doc, 'Draft date: ', 'May 2026')
add_meta_line(doc, 'Source documents: ', 'Current-State Process Documentation; Gap Analysis & Risk Assessment; SOW_Excel_Fire.docx; EFP-Site-Visit-May-2026 task list')
add_meta_line(doc, 'Draft status: ', 'Scaffold — includes roadmap structure, dependency model, and three fully worked sample entries. Remaining findings are intentionally templated pending Justin review and/or source-document rehydration.')

doc.add_heading('1. Purpose & Scope', level=1)
doc.add_paragraph(
    'This Modernization Plan converts the diagnostic findings in the Gap Analysis & Risk Assessment into an action-oriented roadmap. It does not re-diagnose the business or replace the Current-State Process Documentation. The gap analysis remains the dated assessment snapshot; this plan is the moving execution document used to sequence remediation, assign owners, define success criteria, and track milestones.'
)
doc.add_paragraph(
    'The plan is organized around the same workstreams used in the current-state and gap-analysis deliverables: Financial Operations & Controls; Revenue Generation & Customer Management; Project Delivery & Operations; People & Administration; Institutional Knowledge & Key Relationships; Regulatory & Compliance; and Transaction-Specific Items where applicable.'
)

doc.add_heading('2. Roadmap Overview', level=1)
doc.add_paragraph(
    'The roadmap groups remediation work by horizon. The horizon is not simply a calendar label: it indicates when the work can be responsibly started based on deal mechanics, dependency gates, staff capacity, and whether remediation requires new owner authority after close.'
)
rows = [
    ('#1 — Banking signature-stamp / spending authority', 'Pre-close / close', 'Deal Mechanics', 'Justin / Jaclyn / legal / bank', 'Deal close and new banking relationships', 'Resolve through shareholders or operating agreement, bank signature authority, and spending tiers.'),
    ('#2 — Missing I-9/W-4/MI-W4 employee records', 'Pre-close', 'Scheduled', 'Justin / HR agent / Kevin / Betty', 'On-site records review', 'Collect, audit, and retain required employee paperwork; create repeatable onboarding evidence trail.'),
    ('#6 — Inspection delivery concentration', '0–90 days', 'Sequenced', 'Justin / Kevin / Keith / hiring owner', 'Inspection-delivery hire decision', 'Assign NFPA 25 inspection delivery capacity to Konner or a new hire; unlocks multi-year agreements and inspection growth work.'),
    ('#8 — Multi-year inspection agreements', '0–90 days / Year 1', 'Sequenced', 'Justin / inspection owner', 'Inspection delivery capacity', 'Move after delivery capacity is clear; standardize agreement templates and renewal cadence.'),
    ('#9 — Internal controls / segregation of duties', 'Close / 0–90 days', 'Deal Mechanics', 'Justin / Jaclyn / Kevin / Betty', 'Ownership transition and banking changes', 'Replace legacy informal controls with authority matrix, approval workflow, and bank/accounting permissions.'),
    ('#11 — Estimating methodology not documented', 'Pre-close / 0–90 days', 'Scheduled', 'Justin / Kevin / Keith / Scott input', 'May 11 and May 18 estimating sessions', 'Document estimating process sufficiently for someone other than Kevin or Keith to follow.'),
    ('#12 — Declining inspection volume', 'Year 1', 'Sequenced / Open', 'Justin / inspection owner', 'Inspection-delivery hire decision', 'Rebuild inspection pipeline after delivery owner is chosen; likely tied to agreement renewals and declined inbound follow-up.'),
    ('#14 — QA/QC program absent', '0–90 days / Year 1', 'Open', 'TBD', 'Owner and scope decision needed', 'Define inspection/project QA standard, review cadence, and evidence retained.'),
    ('#15 — Declined inspection-adjacent inbound', 'Year 1', 'Sequenced', 'Justin / inspection owner', 'Inspection-delivery hire decision', 'Create intake/referral/callback path once capacity exists.'),
    ('#17 — Field workforce recruiting & retention process', '0–90 days', 'In Flight', 'Justin / StangDS / HR agent', 'Notion process page pending', 'Document StangDS posting process and EFP Interviews funnel; publish human-readable recruiting process.'),
    ('#18 — Field leadership succession / Keith dependency', 'Pre-close / 0–90 days', 'Scheduled / Sequenced', 'Justin / Kevin / Keith', 'Week of May 11 conversations; hiring path', 'Clarify near-term field leadership coverage and medium-term successor development.'),
]
add_status_table(doc, rows)
doc.add_paragraph(
    'Roadmap summary: pre-close work should prioritize legal/compliance records, estimating documentation, field leadership conversations, and deal-mechanics controls. The first 90 days should convert the most fragile owner-dependent routines into documented processes and assign durable owners. Year 1 work should focus on revenue-side segmentation, inspection growth, payroll/timekeeping migration, QA/QC, budgeting, and IT/security hardening. Strategic items should remain explicitly parked unless they support the core acquisition thesis or remove a dependency from the operating company.'
)

doc.add_heading('3. Sequencing & Dependencies', level=1)
doc.add_paragraph(
    'The plan should be managed by dependency chain, not by a flat list of findings. Several findings cannot be solved cleanly until a related owner, hire, or transaction event is resolved.'
)
doc.add_heading('3.1 Critical dependency chains', level=2)
for text in [
    'Inspection-delivery hire decision unlocks Finding #6, Finding #8, Finding #12, Finding #15, and part of Finding #18. Until Excel decides whether Konner, a new NFPA 25-capable hire, or another resource will own inspection delivery, inspection-growth remediation should stay sequenced rather than over-engineered.',
    'Deal close unlocks Finding #1 and Finding #9. The signature-stamp practice and informal control environment are serious, but the durable fix is tied to new ownership, new banking relationships, and shareholders/operating agreement authority tiers.',
    'Estimating documentation is on the near-term critical path because it directly reduces Kevin/Keith dependency and supports continuity of quoted work. The success test is whether the process can be followed without Kevin or Keith’s assistance.',
    'Payroll and timekeeping modernization should follow the QuickBooks Desktop to QuickBooks Online path with QuickBooks Payroll and a practical digital timecard capture method. Early-state digitization can be as simple as crews texting photos of paper timesheets, provided retention and approval evidence are reliable.',
]:
    doc.add_paragraph(text, style='List Bullet')

doc.add_heading('3.2 Critical path items', level=2)
for text in [
    'Pre-close compliance record audit: I-9, W-4, MI-W4, UIA, BFS Act 144, and any employee-license/training evidence that must be retained.',
    'Week of May 11 field leadership and inspection-capacity conversations with Kevin, Keith, Justin, and candidates in the EFP Interviews funnel.',
    'Weeks of May 11 and May 18 estimating process documentation, including Scott’s larger-shop estimating method as a comparison point.',
    'Deal-mechanics controls: authority tiers, check-signing/banking permissions, approval thresholds, and accounting access after close.',
    'QA/QC scoping decision: whether the first system is inspection-focused, project-focused, or both.'
]:
    doc.add_paragraph(text, style='List Bullet')


doc.add_heading('4. Resource & Effort Summary', level=1)
doc.add_paragraph(
    'This section is intended to keep the plan honest about capacity. The highest-value remediation is not the longest list of improvements; it is the smallest sequence of actions that removes key-person risk, closes legal/compliance gaps, and creates enough operating visibility for Justin and Jaclyn to manage the company after close.'
)
resource_rows = [
    ('Justin / Predictive Lines', 'Roadmap owner; process documentation; interviews; operating-design decisions; vendor coordination.', 'Highest constraint. Protect for critical path and owner-only decisions.'),
    ('Jaclyn', 'Financial controls, accounting visibility, budgeting/forecasting, approval authority, payroll/accounting modernization input.', 'Use for finance/control decisions rather than broad project administration.'),
    ('Kevin', 'Source knowledge for estimating, customers, field leadership, vendor/customer relationships, safety history.', 'High leverage but time-limited; use structured extraction sessions.'),
    ('Keith', 'Field operations, estimating assumptions, inspection delivery, QA/QC practices, crew management.', 'Key-person risk; documentation sessions should be concrete and process-tested.'),
    ('Betty / AT&C', 'Payroll, records, accounting process, UIA/tax filings, transition history.', 'Use for source-data extraction and verification; avoid assigning modernization ownership without explicit buy-in.'),
    ('External vendors', 'StangDS, Gauthier/cyber insurance, IT/MDM, bank/legal, potential Scott estimating visit.', 'Use where specialized capacity beats building from scratch.'),
]
add_status_table(doc, [(a, b, '', c, '', '') for a, b, c in resource_rows])

doc.add_heading('5. Workstream-Aligned Remediation Plans', level=1)
doc.add_paragraph(
    'Each entry below should be short enough to manage but specific enough to test. The goal is not prose completeness; the goal is operational closure.'
)

doc.add_heading('5.1 People & Administration', level=2)
add_entry(
    doc, '2', 'Missing I-9/W-4/MI-W4 Employee Records', 'Critical', 'Gap Analysis §3.4; Current-State §2.4',
    'Scheduled — on-site records review and employee-file cleanup identified for the May 11 site-visit workstream.',
    'Justin / HR agent, with Kevin and Betty supplying existing records and employee roster confirmation.',
    'Pre-close compliance item. Predecessors: current employee roster and physical/digital record access. Successors: repeatable onboarding checklist, termination checklist, and HR evidence-retention process.',
    'For every active employee, Excel has either retained the required I-9, W-4, and MI-W4 documentation or a documented remediation note identifying the missing item, responsible party, and completion date. Going forward, the onboarding checklist prevents a worker from reaching “paperwork complete” without these records.',
    'Small to medium. Most effort is records retrieval and exception handling; documentation update is straightforward once gaps are known.',
    'May 11–15: inventory existing files and identify gaps. Week of May 18: collect missing paperwork or document remediation path. 0–90 days: fold into New Employee Onboarding Checklist and HR Compliance Guide.',
    'Confirm who is authorized to request corrected paperwork; confirm where final employee-file records will live; confirm whether historical gaps require counsel/accountant review.'
)

add_template_entry(doc, '17', 'Field Workforce Recruiting & Retention Process Documentation', 'High', 'Gap Analysis §3.4', '0–90 days', 'In Flight', 'Justin / StangDS / HR agent')

doc.add_heading('5.2 Revenue Generation & Customer Management', level=2)
add_template_entry(doc, '8', 'Multi-Year Inspection Agreements', 'High', 'Gap Analysis §3.2', '0–90 days / Year 1', 'Sequenced', 'Justin / inspection owner')
add_template_entry(doc, '12', 'Declining Inspection Volume', 'High', 'Gap Analysis §3.2', 'Year 1', 'Sequenced / Open', 'Justin / inspection owner')
add_template_entry(doc, '15', 'Declined Inspection-Adjacent Inbound', 'High', 'Gap Analysis §3.2', 'Year 1', 'Sequenced', 'Justin / inspection owner')


doc.add_heading('5.3 Project Delivery & Operations', level=2)
add_entry(
    doc, '11', 'Estimating Methodology Not Documented', 'High', 'Gap Analysis §3.3; Current-State §2.2',
    'Scheduled — focus weeks of May 11 and May 18; includes Kevin/Keith working sessions and Scott comparison visit if scheduled.',
    'Justin, with Kevin and Keith as source experts; Scott as external comparison input if the Wisconsin visit occurs.',
    'Near-term key-person-risk reduction. Predecessors: access to recent estimates, rate/cost-buildup assumptions, and Kevin/Keith availability. Successors: estimating training, review checklist, and eventual delegation to a non-owner estimator or operations lead.',
    'A documented estimating process can be followed without Kevin or Keith’s assistance. A test user can take a representative bid opportunity, identify required inputs, apply labor/material/subcontractor/equipment assumptions, produce a quote package, and explain review/approval checkpoints.',
    'Medium. Requires structured interviews and process testing, not just transcription. Expect multiple working sessions and one validation run against a real or historical estimate.',
    'Week of May 11: capture current Kevin/Keith methodology and open questions. Week of May 18: draft v0 and validate against at least one estimate. 0–90 days: convert to maintained procedure with sample estimate package and approval checklist.',
    'Confirm whether Excel uses any stable labor-unit assumptions, historical job-cost feedback, vendor quote standards, or margin targets. Confirm how Scott’s method should influence Excel’s process versus simply serving as a benchmark.'
)
add_entry(
    doc, '14', 'QA/QC Program Absent', 'High', 'Gap Analysis §3.3; Current-State §3 summary',
    'Open — no modernization workstream assigned yet.',
    'TBD. Candidate owners include Justin for design, Keith or future field leader for execution, and inspection/project delivery owner depending on scope.',
    'Decision needed before execution. Predecessors: scope decision on whether QA/QC starts with inspections, installation projects, service work, or all field delivery. Successors: QA checklists, photo/evidence retention, corrective-action loop, and training cadence.',
    'Excel has a defined QA/QC standard with evidence retained. A sample job or inspection can be reviewed after completion and show who checked it, what standard was used, what evidence was captured, and how exceptions were corrected.',
    'Medium to large depending on scope. A narrow inspection-first QA process is much smaller than a full project-delivery QA program.',
    'Decision trigger: assign owner and first-scope by end of 0–90 day planning. After decision: draft checklist, pilot on 2–3 jobs/inspections, revise, then incorporate into field operating rhythm.',
    'What is the first QA/QC surface: NFPA 25 inspections, installation closeout, service calls, or all three? Who has authority to stop/rework field output? What evidence is useful without creating paperwork theater?'
)


doc.add_heading('5.4 Financial Operations & Controls', level=2)
add_template_entry(doc, '1', 'Banking Signature-Stamp Practice / Spending Authority', 'Critical', 'Gap Analysis §3.1', 'Pre-close / close', 'Deal Mechanics', 'Justin / Jaclyn / legal / bank')
add_template_entry(doc, '9', 'Internal Controls & Segregation of Duties', 'Critical', 'Gap Analysis §3.1', 'Close / 0–90 days', 'Deal Mechanics', 'Justin / Jaclyn / Kevin / Betty')
add_template_entry(doc, '[TBD]', 'Payroll and Timekeeping Modernization', 'High', 'Gap Analysis §3.1 / §3.4', '0–90 days / Year 1', 'In Flight / Sequenced', 'Justin / Jaclyn / AT&C / payroll vendor')


doc.add_heading('5.5 Institutional Knowledge & Key Relationships', level=2)
add_template_entry(doc, '18', 'Field Leadership Succession / Keith Dependency', 'Critical / High', 'Gap Analysis §4', 'Pre-close / 0–90 days', 'Scheduled / Sequenced', 'Justin / Kevin / Keith')
add_template_entry(doc, '6', 'Inspection Delivery Concentration', 'Critical / High', 'Gap Analysis §3.3 / §4', '0–90 days', 'Sequenced', 'Justin / Kevin / Keith / hiring owner')


doc.add_heading('5.6 Regulatory & Compliance', level=2)
add_template_entry(doc, '[TBD]', 'BFS Act 144 Certification Verification', 'High', 'Gap Analysis §3.6; Current-State §2.6', 'Pre-close', 'Scheduled', 'Justin / Kevin')
add_template_entry(doc, '[TBD]', 'Michigan UIA Filing Cleanup', 'High', 'Gap Analysis §3.6', 'Pre-close / 0–90 days', 'Scheduled / In Flight', 'AT&C / Justin')
add_template_entry(doc, '[TBD]', 'Safety Program Documentation and Calendar', 'High', 'Gap Analysis §3.6', '0–90 days', 'In Flight', 'HR agent / Justin / Kevin')


doc.add_heading('6. Items Not Yet Scoped', level=1)
doc.add_paragraph('The following items are intentionally not converted into detailed workplans until a decision trigger is met. This keeps the modernization plan from pretending to have ownership where no owner has been named.')
not_scoped = [
    ('QA/QC program', 'Assign first-scope and operating owner.'),
    ('ISN compliance calendar', 'Decide whether customer-specific compliance tracking sits in HR, operations, or customer management.'),
    ('Customer-specific compliance tracking', 'Inventory customer portals/requirements and assign maintenance owner.'),
    ('Fleet maintenance', 'Confirm current maintenance records and whether fleet process belongs with field operations or finance/admin.'),
    ('COI tracking', 'Choose tracking system and renewal owner.'),
    ('Budgeting / forecasting', 'Wait for post-close accounting visibility and Jaclyn/Justin operating cadence.'),
    ('MDM / device management', 'Decide company-device policy after M365 and field communication architecture stabilize.'),
    ('Chart of accounts cleanup', 'Sequence with QuickBooks Online migration and reporting design.'),
    ('Revenue-side segmentation', 'Sequence with customer/job database maturity and inspection/service strategy.'),
    ('Stale AR', 'Requires AR aging review and customer-specific collection decisions.'),
    ('Inventory', 'Requires physical inventory practice decision and accounting treatment.'),
]
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
set_cell_text(t.rows[0].cells[0], 'Item', True); set_cell_shading(t.rows[0].cells[0], 'D9EAF7')
set_cell_text(t.rows[0].cells[1], 'Decision Trigger', True); set_cell_shading(t.rows[0].cells[1], 'D9EAF7')
for item, trigger in not_scoped:
    cells = t.add_row().cells
    set_cell_text(cells[0], item)
    set_cell_text(cells[1], trigger)


doc.add_heading('7. Out-of-Scope Items', level=1)
doc.add_paragraph(
    'The following work may be valuable but should not be treated as gap-analysis remediation unless and until Justin explicitly brings it into scope. Keeping this boundary clean prevents the modernization plan from absorbing every attractive adjacent idea.'
)
for text in [
    'New product-line investigations, including portable fire extinguisher service or fire alarm inspection expansion, unless tied to the acquisition thesis and separately scoped.',
    'Marquette warehouse or real-estate options beyond what is required for operational continuity.',
    'Sister-company or non-union entity design, except as a separate strategic project with legal/tax review.',
    'Broad brand/marketing redesign not directly required to address inspection volume, recruiting, or customer-management gaps.'
]:
    doc.add_paragraph(text, style='List Bullet')


doc.add_heading('Appendix A — Source-Document Update Notes', level=1)
doc.add_paragraph(
    'This scaffold was prepared from the May 2026 rehydration brief because the original ai-team source folder was not mounted in the OpenClaw host at build time. Before finalizing the engagement packet, update the current-state and gap-analysis documents directly from their source DOCX files and reconcile any finding numbers/titles against the actual 44-finding table.'
)
doc.add_paragraph('Current-State Process Documentation updates to confirm: estimating rates/cost-buildup, BFS Act 144 certification verification, QA/QC verification, and First Bank LOC limit.', style='List Bullet')
doc.add_paragraph('Gap Analysis updates to confirm: remediation status text remains paired to the correct finding; Certified Payroll / Banking Signature-Stamp mismatch bug does not recur; summary table status column matches per-finding remediation tags.', style='List Bullet')
doc.add_paragraph('Modernization Plan updates to confirm: all 44 findings are represented or deliberately excluded; roadmap table uses final finding titles and cross-references from the current gap-analysis draft.', style='List Bullet')

doc.save(OUT)
print(OUT)
